# Pythia
# Copyright (c) 2025 Kevin Wyjad
# Licensed under the Pythia Non-Commercial Public License v1.0.
# See the LICENSE file in the project root for details.

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping, Optional, Sequence, Tuple

from docx import Document
from resolver.db import duckdb_io

from pythia.db.schema import get_db_url


def _safe_load_json(raw: object, default: Any) -> Any:
    if raw is None:
        return default

    if isinstance(raw, (dict, list)):
        return raw

    try:
        return json.loads(raw)
    except Exception:
        return default


def _connect() -> Any:
    db_url = get_db_url()
    return duckdb_io.get_db(db_url)


def _load_question(
    con: Any, question_id: str
) -> Tuple[Mapping[str, Any], List[Mapping[str, Any]], Optional[Mapping[str, Any]]]:
    row = con.execute(
        """
        SELECT question_id, hs_run_id, scenario_ids_json, iso3, hazard_code, metric,
               target_month, window_start_date, window_end_date, wording, status,
               pythia_metadata_json
        FROM questions
        WHERE question_id = ?
        LIMIT 1
        """,
        [question_id],
    ).fetchone()

    if not row:
        raise ValueError(f"Question {question_id} not found in questions table.")

    columns = [
        "question_id",
        "hs_run_id",
        "scenario_ids_json",
        "iso3",
        "hazard_code",
        "metric",
        "target_month",
        "window_start_date",
        "window_end_date",
        "wording",
        "status",
        "pythia_metadata_json",
    ]
    question = {col: value for col, value in zip(columns, row)}

    hs_run_id = question.get("hs_run_id") or ""
    iso3 = (question.get("iso3") or "").upper()
    scenario_ids: Sequence[str] = _safe_load_json(question.get("scenario_ids_json"), [])

    scenarios: List[Mapping[str, Any]] = []
    if hs_run_id and scenario_ids:
        placeholders = ",".join(["?"] * len(scenario_ids))
        scenario_rows = con.execute(
            f"""
            SELECT scenario_id, iso3, hazard_code, scenario_title, likely_month,
                   probability_pct, pin_best_guess, pa_best_guess, scenario_markdown
            FROM hs_scenarios
            WHERE hs_run_id = ? AND scenario_id IN ({placeholders})
            ORDER BY iso3, scenario_title, scenario_id
            """,
            [hs_run_id, *scenario_ids],
        ).fetchall()

        scenario_columns = [
            "scenario_id",
            "iso3",
            "hazard_code",
            "scenario_title",
            "likely_month",
            "probability_pct",
            "pin_best_guess",
            "pa_best_guess",
            "scenario_markdown",
        ]
        scenarios = [dict(zip(scenario_columns, s)) for s in scenario_rows]

    country_report: Optional[Mapping[str, Any]] = None
    if hs_run_id and iso3:
        report_row = con.execute(
            """
            SELECT report_markdown, sources_json
            FROM hs_country_reports
            WHERE hs_run_id = ? AND UPPER(iso3) = ?
            LIMIT 1
            """,
            [hs_run_id, iso3],
        ).fetchone()
        if report_row:
            country_report = {
                "report_markdown": report_row[0],
                "sources_json": report_row[1],
            }

    return question, scenarios, country_report


def _resolve_run_id(con: Any, question_id: str, explicit_run_id: Optional[str]) -> Optional[str]:
    if explicit_run_id:
        return explicit_run_id

    row = con.execute(
        """
        SELECT run_id
        FROM forecasts_ensemble
        WHERE question_id = ?
        ORDER BY created_at DESC
        LIMIT 1
        """,
        [question_id],
    ).fetchone()

    if not row:
        return None

    return row[0]


def _load_spd_forecasts(
    con: Any, question_id: str, run_id: str
) -> Tuple[Dict[int, List[float]], Dict[int, Optional[float]], List[Mapping[str, Any]]]:
    ensemble_rows = con.execute(
        """
        SELECT month_index, bucket_index, probability, ev_value
        FROM forecasts_ensemble
        WHERE question_id = ? AND run_id = ?
        ORDER BY month_index, bucket_index
        """,
        [question_id, run_id],
    ).fetchall()

    ensemble_probs: Dict[int, List[float]] = {}
    ensemble_ev: Dict[int, Optional[float]] = {}

    for month_index, bucket_index, probability, ev_value in ensemble_rows:
        month_idx = int(month_index)
        bucket_idx = int(bucket_index)
        probs = ensemble_probs.setdefault(month_idx, [0.0] * 5)
        probs[bucket_idx - 1] = float(probability)
        if ev_value is not None:
            ensemble_ev.setdefault(month_idx, float(ev_value))

    raw_rows = con.execute(
        """
        SELECT model_name, month_index, bucket_index, probability,
               ok, elapsed_ms, cost_usd, prompt_tokens, completion_tokens, total_tokens
        FROM forecasts_raw
        WHERE question_id = ? AND run_id = ?
        ORDER BY model_name, month_index, bucket_index
        """,
        [question_id, run_id],
    ).fetchall()

    raw_models: Dict[str, Dict[str, Any]] = {}
    for (
        model_name,
        month_index,
        bucket_index,
        probability,
        ok,
        elapsed_ms,
        cost_usd,
        prompt_tokens,
        completion_tokens,
        total_tokens,
    ) in raw_rows:
        model_key = str(model_name)
        model_entry = raw_models.setdefault(
            model_key,
            {
                "model_name": model_key,
                "ok": bool(ok),
                "elapsed_ms": int(elapsed_ms or 0),
                "cost_usd": float(cost_usd or 0.0),
                "prompt_tokens": int(prompt_tokens or 0),
                "completion_tokens": int(completion_tokens or 0),
                "total_tokens": int(total_tokens or 0),
                "spd": {},
            },
        )

        month_idx = int(month_index)
        bucket_idx = int(bucket_index)
        probs = model_entry["spd"].setdefault(month_idx, [0.0] * 5)
        probs[bucket_idx - 1] = float(probability)

    sorted_models = [raw_models[key] for key in sorted(raw_models.keys())]
    return ensemble_probs, ensemble_ev, sorted_models


def _load_llm_calls(con: Any, question_id: str, run_id: str) -> Dict[str, Any]:
    columns = [
        "call_id",
        "call_type",
        "model_name",
        "provider",
        "model_id",
        "prompt_text",
        "response_text",
        "elapsed_ms",
        "prompt_tokens",
        "completion_tokens",
        "total_tokens",
        "cost_usd",
        "error_text",
        "timestamp",
    ]

    rows = con.execute(
        """
        SELECT call_id, call_type, model_name, provider, model_id,
               prompt_text, response_text, elapsed_ms, prompt_tokens,
               completion_tokens, total_tokens, cost_usd, error_text, timestamp
        FROM llm_calls
        WHERE question_id = ? AND run_id = ? AND call_type IN ('research', 'forecast')
        ORDER BY timestamp, call_id
        """,
        [question_id, run_id],
    ).fetchall()

    research_calls: List[Mapping[str, Any]] = []
    forecast_calls: List[Mapping[str, Any]] = []

    for row in rows:
        row_dict = {col: value for col, value in zip(columns, row)}
        if row_dict.get("call_type") == "research":
            research_calls.append(row_dict)
        else:
            forecast_calls.append(row_dict)

    return {"research_calls": research_calls, "forecast_calls": forecast_calls}


def _load_question_context(con: Any, question_id: str, run_id: str) -> Mapping[str, Any]:
    row = con.execute(
        """
        SELECT iso3, hazard_code, metric,
               snapshot_start_month, snapshot_end_month,
               pa_history_json, context_json
        FROM question_context
        WHERE question_id = ? AND run_id = ?
        LIMIT 1
        """,
        [question_id, run_id],
    ).fetchone()

    if not row:
        return {}

    columns = [
        "iso3",
        "hazard_code",
        "metric",
        "snapshot_start_month",
        "snapshot_end_month",
        "pa_history_json",
        "context_json",
    ]
    context = {col: value for col, value in zip(columns, row)}
    context["pa_history"] = _safe_load_json(context.get("pa_history_json"), [])
    context["context"] = _safe_load_json(context.get("context_json"), {})
    return context


def _load_calibration(con: Any, hazard_code: str, metric: str) -> Mapping[str, Any]:
    weights_rows = con.execute(
        """
        SELECT model_name, weight
        FROM calibration_weights
        WHERE hazard_code = ? AND metric = ?
        ORDER BY model_name
        """,
        [hazard_code, metric],
    ).fetchall()

    weights = [
        {"model_name": row[0], "weight": float(row[1])}
        for row in weights_rows
    ]

    advice_row = con.execute(
        """
        SELECT advice_text
        FROM calibration_advice
        WHERE hazard_code = ? AND metric = ?
        ORDER BY updated_at DESC
        LIMIT 1
        """,
        [hazard_code, metric],
    ).fetchone()

    advice_text = advice_row[0] if advice_row else ""
    return {"weights": weights, "advice_text": advice_text or ""}


def _add_scenarios(doc: Document, scenarios: Iterable[Mapping[str, Any]]) -> None:
    scenarios_list = list(scenarios)
    if not scenarios_list:
        doc.add_paragraph("No HS scenarios found for this question.")
        return

    doc.add_paragraph("Key scenarios:")
    table = doc.add_table(rows=1, cols=5)
    header = table.rows[0].cells
    header[0].text = "Scenario Title"
    header[1].text = "Hazard"
    header[2].text = "Likely Month"
    header[3].text = "Probability (%)"
    header[4].text = "PA Best Guess"

    for scenario in scenarios_list:
        row = table.add_row().cells
        row[0].text = str(scenario.get("scenario_title") or "")
        row[1].text = str(scenario.get("hazard_code") or "")
        row[2].text = str(scenario.get("likely_month") or "")
        row[3].text = str(scenario.get("probability_pct") or "")
        row[4].text = str(scenario.get("pa_best_guess") or "")


def _add_country_report_excerpt(doc: Document, country_report: Optional[Mapping[str, Any]]) -> None:
    if not country_report:
        return

    markdown = country_report.get("report_markdown") or ""
    if not markdown:
        return

    doc.add_paragraph()
    doc.add_paragraph("Country report excerpt:", style="Intense Quote")
    lines = markdown.splitlines()
    excerpt = "\n".join(lines[:15])
    doc.add_paragraph(excerpt)


def _add_ensemble_table(doc: Document, ensemble_probs: Mapping[int, Sequence[float]]) -> None:
    if not ensemble_probs:
        doc.add_paragraph("No ensemble SPD data found.")
        return

    doc.add_paragraph("Ensemble SPD (months × buckets):")
    max_month = max(ensemble_probs.keys())
    table = doc.add_table(rows=max_month + 1, cols=6)
    header = table.rows[0].cells
    header[0].text = "Month"
    for idx in range(5):
        header[idx + 1].text = f"Bucket {idx + 1}"

    for month in range(1, max_month + 1):
        row = table.rows[month].cells
        row[0].text = str(month)
        probs = ensemble_probs.get(month) or [0.0] * 5
        for idx, prob in enumerate(probs):
            row[idx + 1].text = f"{float(prob):.3f}"


def _add_ev_table(doc: Document, ensemble_ev: Mapping[int, Optional[float]]) -> None:
    if not ensemble_ev:
        doc.add_paragraph("No expected PA values found for ensemble.")
        return

    doc.add_paragraph("Expected people affected per month (ensemble):")
    table = doc.add_table(rows=1 + len(ensemble_ev), cols=2)
    header = table.rows[0].cells
    header[0].text = "Month"
    header[1].text = "Expected PA"

    for month in sorted(ensemble_ev.keys()):
        row = table.add_row().cells
        row[0].text = str(month)
        row[1].text = f"{float(ensemble_ev[month]):,.0f}"


def _add_model_summary(doc: Document, raw_models: Sequence[Mapping[str, Any]]) -> None:
    doc.add_paragraph("Per-model SPD summary:")
    if not raw_models:
        doc.add_paragraph("No per-model SPD data found.")
        return

    table = doc.add_table(rows=1, cols=6)
    header = table.rows[0].cells
    header[0].text = "Model"
    header[1].text = "OK"
    header[2].text = "Time (ms)"
    header[3].text = "Cost (USD)"
    header[4].text = "Tokens"
    header[5].text = "month_1 (first 3 buckets)"

    for model in raw_models:
        row = table.add_row().cells
        row[0].text = str(model.get("model_name") or "")
        row[1].text = "yes" if model.get("ok") else "no"
        row[2].text = str(model.get("elapsed_ms") or 0)
        row[3].text = f"{float(model.get('cost_usd') or 0.0):.4f}"
        row[4].text = str(model.get("total_tokens") or 0)
        spd = model.get("spd") or {}
        month1_probs = spd.get(1, [0.0] * 5)
        slice_text = ", ".join(f"{prob:.2f}" for prob in month1_probs[:3])
        row[5].text = slice_text


def _add_resolver_context(doc: Document, context: Mapping[str, Any]) -> None:
    if not context:
        doc.add_paragraph("No Resolver context found for this question/run.")
        return

    doc.add_paragraph(
        f"Snapshot: {context.get('snapshot_start_month', '')} → {context.get('snapshot_end_month', '')}"
    )
    history = context.get("pa_history") or []
    if not history:
        doc.add_paragraph("No PA history found.")
        return

    table = doc.add_table(rows=1, cols=3)
    header = table.rows[0].cells
    header[0].text = "Month"
    header[1].text = "PA"
    header[2].text = "Source"

    for item in history:
        row = table.add_row().cells
        row[0].text = str(item.get("ym") or "")
        row[1].text = str(item.get("value") or "")
        row[2].text = str(item.get("source") or "")


def _add_calibration(doc: Document, calibration: Mapping[str, Any]) -> None:
    weights = calibration.get("weights") or []
    advice_text = calibration.get("advice_text") or ""

    if weights:
        doc.add_paragraph("Calibration weights (per model):")
        table = doc.add_table(rows=1, cols=2)
        header = table.rows[0].cells
        header[0].text = "Model"
        header[1].text = "Weight"
        for weight in weights:
            row = table.add_row().cells
            row[0].text = str(weight.get("model_name") or "")
            row[1].text = f"{float(weight.get('weight') or 0.0):.3f}"
    else:
        doc.add_paragraph("No calibration weights found for this hazard/metric.")

    if advice_text:
        doc.add_paragraph("Calibration advice:")
        doc.add_paragraph(advice_text)


def _add_llm_research(doc: Document, research_calls: Sequence[Mapping[str, Any]]) -> None:
    if not research_calls:
        doc.add_paragraph("No research calls found for this question/run.")
        return

    last_call = research_calls[-1]
    provider = last_call.get("provider") or ""
    model_id = last_call.get("model_id") or ""
    model_name = last_call.get("model_name") or ""
    cost = float(last_call.get("cost_usd") or 0.0)
    tokens = int(last_call.get("total_tokens") or 0)

    doc.add_paragraph(f"Research LLM: {provider}/{model_id} ({model_name})")
    doc.add_paragraph(f"Research usage: tokens={tokens}, cost=${cost:.4f}")
    doc.add_paragraph("Research prompt (truncated):")
    doc.add_paragraph((last_call.get("prompt_text") or "")[:1000])
    doc.add_paragraph("Research response (truncated):")
    doc.add_paragraph((last_call.get("response_text") or "")[:1000])


def _add_forecast_annex(doc: Document, forecast_calls: Sequence[Mapping[str, Any]]) -> None:
    if not forecast_calls:
        doc.add_paragraph("No forecast LLM calls logged for this question/run.")
        return

    for call in forecast_calls:
        doc.add_heading(str(call.get("model_name") or ""), level=2)
        doc.add_paragraph(
            f"Model: {call.get('provider')}/{call.get('model_id')} ({call.get('model_name')}), "
            f"time={call.get('elapsed_ms')} ms, tokens={call.get('total_tokens')}, cost=${float(call.get('cost_usd') or 0.0):.4f}"
        )
        doc.add_paragraph("Prompt (truncated):")
        doc.add_paragraph((call.get("prompt_text") or "")[:2000])
        doc.add_paragraph("Response (truncated):")
        doc.add_paragraph((call.get("response_text") or "")[:2000])


def _build_document(
    output_path: Path,
    question: Mapping[str, Any],
    scenarios: Sequence[Mapping[str, Any]],
    country_report: Optional[Mapping[str, Any]],
    ensemble_probs: Mapping[int, Sequence[float]],
    ensemble_ev: Mapping[int, Optional[float]],
    raw_models: Sequence[Mapping[str, Any]],
    llm_info: Mapping[str, Any],
    context: Mapping[str, Any],
    calibration: Mapping[str, Any],
) -> None:
    doc = Document()

    question_id = question.get("question_id") or ""
    wording = question.get("wording") or ""
    iso3 = (question.get("iso3") or "").upper()
    hazard_code = (question.get("hazard_code") or "").upper()
    metric = question.get("metric") or ""
    hs_run_id = question.get("hs_run_id") or ""
    window_start = question.get("window_start_date") or ""
    window_end = question.get("window_end_date") or ""

    doc.add_heading("Pythia Forecast Question Report", level=0)
    doc.add_paragraph(f"Question ID: {question_id}")
    doc.add_paragraph(f"HS Run ID: {hs_run_id}")

    doc.add_heading("1. Question Summary", level=1)
    doc.add_paragraph(wording)
    summary = doc.add_paragraph()
    summary.add_run("Country: ").bold = True
    summary.add_run(iso3)
    summary.add_run("   Hazard: ").bold = True
    summary.add_run(hazard_code)
    summary.add_run("   Metric: ").bold = True
    summary.add_run(metric)

    if window_start and window_end:
        doc.add_paragraph(f"Forecast window: {window_start} → {window_end}")

    scenario_titles = [s.get("scenario_title") for s in scenarios if s.get("scenario_title")]
    if scenario_titles:
        doc.add_paragraph("HS scenarios linked to this question:")
        for title in scenario_titles:
            doc.add_paragraph(f"- {title}", style="List Bullet")

    doc.add_heading("2. Horizon Scanner Context", level=1)
    _add_scenarios(doc, scenarios)
    _add_country_report_excerpt(doc, country_report)

    doc.add_heading("3. Research", level=1)
    _add_llm_research(doc, llm_info.get("research_calls") or [])

    doc.add_heading("4. Forecast (SPD Ensemble)", level=1)
    _add_ensemble_table(doc, ensemble_probs)
    _add_ev_table(doc, ensemble_ev)
    doc.add_paragraph()
    _add_model_summary(doc, raw_models)

    doc.add_heading("5. Resolver Context (36-month snapshot)", level=1)
    _add_resolver_context(doc, context)

    doc.add_heading("6. Calibration", level=1)
    _add_calibration(doc, calibration)

    doc.add_heading("Annex: Raw Forecast Explanations (Truncated)", level=1)
    _add_forecast_annex(doc, llm_info.get("forecast_calls") or [])

    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Export a Pythia forecast question to Word.")
    parser.add_argument("--question-id", required=True, help="Question ID to export.")
    parser.add_argument(
        "--run-id",
        required=False,
        help=(
            "Optional forecaster run_id. If omitted, the latest run_id with forecasts_ensemble "
            "rows for this question is used."
        ),
    )
    args = parser.parse_args()

    conn = _connect()
    try:
        question, scenarios, country_report = _load_question(conn, args.question_id)
        run_id = _resolve_run_id(conn, args.question_id, args.run_id)
        if not run_id:
            raise ValueError(
                f"No forecasts_ensemble rows found for question_id={args.question_id}"
            )

        ensemble_probs, ensemble_ev, raw_models = _load_spd_forecasts(
            conn, args.question_id, run_id
        )
        llm_info = _load_llm_calls(conn, args.question_id, run_id)
        context = _load_question_context(conn, args.question_id, run_id)

        hazard_code = question.get("hazard_code") or ""
        metric = question.get("metric") or ""
        calibration = _load_calibration(conn, hazard_code, metric)
    finally:
        duckdb_io.close_db(conn)

    output_name = f"Forecast_Q_{args.question_id}.docx"
    output_path = Path(output_name)
    _build_document(
        output_path,
        question,
        scenarios,
        country_report,
        ensemble_probs,
        ensemble_ev,
        raw_models,
        llm_info,
        context,
        calibration,
    )
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
