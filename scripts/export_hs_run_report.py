from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Iterable, Mapping, Sequence

from docx import Document
from resolver.db import duckdb_io

from pythia.db.schema import ensure_schema, get_db_url


def _load_json(raw: object) -> dict:
    if raw is None:
        return {}

    if isinstance(raw, Mapping):
        return dict(raw)

    try:
        return json.loads(raw)
    except Exception:
        return {}


def _format_probability(prob_pct: float | None, probability_text: str | None) -> str:
    if prob_pct is None:
        return probability_text or ""

    try:
        value = float(prob_pct)
    except Exception:
        return str(prob_pct)

    if value.is_integer():
        return f"{int(value)}%"

    return f"{value:.2f}%"


def _add_formatted_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _add_markdown(document: Document, markdown_text: str) -> None:
    for line in markdown_text.splitlines():
        stripped = line.rstrip()

        if not stripped:
            document.add_paragraph()
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        number_match = re.match(r"^\d+[.)]\s+(.*)$", stripped)

        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            document.add_heading(text, level=level)
            continue

        if bullet_match:
            text = bullet_match.group(1).strip()
            paragraph = document.add_paragraph(style="List Bullet")
            _add_formatted_runs(paragraph, text)
            continue

        if number_match:
            text = number_match.group(1).strip()
            paragraph = document.add_paragraph(style="List Number")
            _add_formatted_runs(paragraph, text)
            continue

        paragraph = document.add_paragraph()
        _add_formatted_runs(paragraph, stripped)


def _fetch_run_metadata(conn, hs_run_id: str) -> Mapping[str, object]:
    row = conn.execute(
        """
        SELECT hs_run_id, generated_at, git_sha, config_profile, countries_json
        FROM hs_runs
        WHERE hs_run_id = ?
        LIMIT 1
        """,
        [hs_run_id],
    ).fetchone()

    if not row:
        raise ValueError(f"No hs_run found for hs_run_id={hs_run_id}")

    hs_run_id_val, generated_at, git_sha, config_profile, countries_json = row
    countries = []
    try:
        countries = sorted(json.loads(countries_json or "[]"))
    except Exception:
        countries = []

    return {
        "hs_run_id": hs_run_id_val,
        "generated_at": generated_at,
        "git_sha": git_sha,
        "config_profile": config_profile,
        "countries": countries,
    }


def _fetch_scenarios(conn, hs_run_id: str) -> list[dict[str, object]]:
    rows = conn.execute(
        """
        SELECT iso3, hazard_code, scenario_title, likely_month, probability_pct,
               pin_best_guess, pa_best_guess, scenario_json
        FROM hs_scenarios
        WHERE hs_run_id = ?
        ORDER BY iso3, scenario_title, hazard_code
        """,
        [hs_run_id],
    ).fetchall()

    scenarios: list[dict[str, object]] = []
    for iso3, hazard_code, scenario_title, likely_month, probability_pct, pin_best_guess, pa_best_guess, scenario_json in rows:
        parsed_json = _load_json(scenario_json)
        country_name = (
            parsed_json.get("country")
            or parsed_json.get("country_name")
            or parsed_json.get("countryName")
        )
        probability_text = parsed_json.get("probability_text") or parsed_json.get("probability")

        scenarios.append(
            {
                "iso3": (iso3 or "").upper(),
                "hazard_code": hazard_code or "",
                "scenario_title": scenario_title or "",
                "likely_month": likely_month or "",
                "probability_pct": probability_pct,
                "probability_text": probability_text,
                "pin_best_guess": pin_best_guess,
                "pa_best_guess": pa_best_guess,
                "country_name": country_name,
                "scenario_json": parsed_json,
            }
        )

    return scenarios


def _fetch_country_reports(conn, hs_run_id: str) -> list[dict[str, str]]:
    rows = conn.execute(
        """
        SELECT iso3, report_markdown
        FROM hs_country_reports
        WHERE hs_run_id = ?
        ORDER BY iso3
        """,
        [hs_run_id],
    ).fetchall()

    reports: list[dict[str, str]] = []
    for iso3, report_markdown in rows:
        reports.append({"iso3": (iso3 or "").upper(), "report_markdown": report_markdown or ""})

    return reports


def _build_country_lookup(scenarios: Iterable[Mapping[str, object]]) -> dict[str, str]:
    lookup: dict[str, str] = {}
    for scenario in scenarios:
        iso3 = str(scenario.get("iso3") or "").upper()
        country_name = scenario.get("country_name")
        if iso3 and isinstance(country_name, str) and country_name:
            lookup.setdefault(iso3, country_name)

    return lookup


def _derive_country_name(iso3: str, lookup: Mapping[str, str], report_markdown: str) -> str:
    if iso3 in lookup and lookup[iso3]:
        return lookup[iso3]

    heading_match = re.search(r"^#\s+(.+)$", report_markdown, flags=re.MULTILINE)
    if heading_match:
        return heading_match.group(1).strip()

    return iso3


def _extract_model_names(scenarios: Sequence[Mapping[str, object]]) -> list[str]:
    models: set[str] = set()
    for scenario in scenarios:
        scenario_json = scenario.get("scenario_json")
        if not isinstance(scenario_json, Mapping):
            continue

        model_value = scenario_json.get("model") or scenario_json.get("model_name")
        if isinstance(model_value, str) and model_value.strip():
            models.add(model_value.strip())
        elif isinstance(model_value, Sequence) and not isinstance(model_value, (str, bytes)):
            for item in model_value:
                if isinstance(item, str) and item.strip():
                    models.add(item.strip())

    if not models:
        return ["Unknown"]

    return sorted(models)


def _build_summary_table(document: Document, scenarios: Sequence[Mapping[str, object]], country_lookup: Mapping[str, str]) -> None:
    heading = document.add_heading("Global Scenario Summary", level=1)
    heading.alignment = 0

    table = document.add_table(rows=len(scenarios) + 1, cols=7)
    headers = [
        "Country",
        "Scenario Title",
        "Hazard",
        "Likely Month",
        "Probability",
        "PIN Best Guess",
        "PA Best Guess",
    ]

    for idx, title in enumerate(headers):
        table.rows[0].cells[idx].text = title

    for row_idx, scenario in enumerate(scenarios, start=1):
        iso3 = str(scenario.get("iso3") or "").upper()
        country_name = country_lookup.get(iso3, iso3)
        probability_pct = scenario.get("probability_pct")
        prob_text = scenario.get("probability_text")
        probability = _format_probability(probability_pct, prob_text if isinstance(prob_text, str) else None)

        table.rows[row_idx].cells[0].text = country_name
        table.rows[row_idx].cells[1].text = str(scenario.get("scenario_title") or "")
        table.rows[row_idx].cells[2].text = str(scenario.get("hazard_code") or "")
        table.rows[row_idx].cells[3].text = str(scenario.get("likely_month") or "")
        table.rows[row_idx].cells[4].text = probability
        table.rows[row_idx].cells[5].text = str(scenario.get("pin_best_guess") or "")
        table.rows[row_idx].cells[6].text = str(scenario.get("pa_best_guess") or "")


def _append_country_reports(
    document: Document,
    reports: Sequence[Mapping[str, str]],
    country_lookup: Mapping[str, str],
) -> None:
    if not reports:
        document.add_paragraph("No country reports were found for this HS run.")
        return

    document.add_page_break()
    document.add_heading("Country Reports", level=1)

    for idx, report in enumerate(reports):
        iso3 = report.get("iso3") or ""
        report_markdown = report.get("report_markdown") or ""
        country_name = _derive_country_name(iso3, country_lookup, report_markdown)

        if idx > 0:
            document.add_page_break()

        document.add_heading(f"{country_name} ({iso3})", level=2)
        _add_markdown(document, report_markdown)


def export_hs_run_report(hs_run_id: str, output_path: Path | None = None) -> Path:
    db_url = get_db_url()
    conn = duckdb_io.get_db(db_url)

    try:
        ensure_schema(conn)

        run_metadata = _fetch_run_metadata(conn, hs_run_id)
        scenarios = _fetch_scenarios(conn, hs_run_id)
        reports = _fetch_country_reports(conn, hs_run_id)
    finally:
        duckdb_io.close_db(conn)

    country_lookup = _build_country_lookup(scenarios)
    models = _extract_model_names(scenarios)

    sorted_scenarios = sorted(
        scenarios,
        key=lambda s: (
            country_lookup.get(str(s.get("iso3") or "").upper(), str(s.get("iso3") or "")),
            str(s.get("scenario_title") or ""),
        ),
    )

    country_order = {iso: idx for idx, iso in enumerate(run_metadata.get("countries", []))}
    sorted_reports = sorted(
        reports,
        key=lambda r: (
            country_order.get(str(r.get("iso3") or ""), len(country_order)),
            str(r.get("iso3") or ""),
        ),
    )

    document = Document()
    document.add_heading("Horizon Scanner Run Report", 0)
    document.add_paragraph(f"Run ID: {run_metadata.get('hs_run_id')}")
    document.add_paragraph(f"Generated At: {run_metadata.get('generated_at')}")
    document.add_paragraph(f"Models: {', '.join(models)}")
    document.add_paragraph(f"Config Profile: {run_metadata.get('config_profile') or 'N/A'}")

    document.add_page_break()
    _build_summary_table(document, sorted_scenarios, country_lookup)

    _append_country_reports(document, sorted_reports, country_lookup)

    target_path = output_path or Path(f"HS_Run_{hs_run_id}.docx")
    target_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(target_path)
    return target_path


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Export a Horizon Scanner run report to Word")
    parser.add_argument("--hs-run-id", required=True, help="HS run identifier to export")
    parser.add_argument(
        "--output",
        help="Optional path for the generated .docx file (defaults to HS_Run_<hs_run_id>.docx)",
    )
    return parser


def main(argv: Sequence[str] | None = None) -> None:
    parser = _build_parser()
    args = parser.parse_args(argv)

    output = export_hs_run_report(args.hs_run_id, Path(args.output) if args.output else None)
    print(f"✅ Exported HS run report to {output}")


if __name__ == "__main__":
    main()

