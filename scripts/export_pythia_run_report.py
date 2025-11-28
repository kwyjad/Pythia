from __future__ import annotations

import argparse
import json
from collections import defaultdict
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import duckdb
from docx import Document

from pythia.db import schema


def _connect_duckdb() -> duckdb.DuckDBPyConnection:
    """Return a read-only DuckDB connection using the configured Pythia URL."""

    con = schema.connect(read_only=True)
    return con


def load_hs_run_summary(
    con: duckdb.DuckDBPyConnection,
    hs_run_id: str,
) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """Return (hs_run_row, hs_scenarios_rows)."""

    hs_row = con.execute(
        """
        SELECT hs_run_id, generated_at, git_sha, config_profile, countries_json
        FROM hs_runs
        WHERE hs_run_id = ?
        """,
        [hs_run_id],
    ).fetchone()
    if not hs_row:
        raise RuntimeError(f"hs_run_id={hs_run_id} not found in hs_runs.")

    hs_cols = ["hs_run_id", "generated_at", "git_sha", "config_profile", "countries_json"]
    hs = dict(zip(hs_cols, hs_row))

    rows = con.execute(
        """
        SELECT scenario_id, iso3, hazard_code, scenario_title, likely_month,
               probability_pct, pin_best_guess, pa_best_guess
        FROM hs_scenarios
        WHERE hs_run_id = ?
        ORDER BY iso3, scenario_title
        """,
        [hs_run_id],
    ).fetchall()
    s_cols = [
        "scenario_id",
        "iso3",
        "hazard_code",
        "scenario_title",
        "likely_month",
        "probability_pct",
        "pin_best_guess",
        "pa_best_guess",
    ]
    scenarios = [dict(zip(s_cols, r)) for r in rows]
    return hs, scenarios


def load_pythia_questions_for_hs(
    con: duckdb.DuckDBPyConnection,
    hs_run_id: str,
) -> List[Dict[str, Any]]:
    """
    Return all Pythia questions linked to this hs_run_id, from questions table.
    """

    rows = con.execute(
        """
        SELECT question_id, iso3, hazard_code, metric,
               target_month, window_start_date, window_end_date,
               wording, status, scenario_ids_json, pythia_metadata_json
        FROM questions
        WHERE hs_run_id = ?
        ORDER BY iso3, hazard_code, question_id
        """,
        [hs_run_id],
    ).fetchall()

    cols = [
        "question_id",
        "iso3",
        "hazard_code",
        "metric",
        "target_month",
        "window_start_date",
        "window_end_date",
        "wording",
        "status",
        "scenario_ids_json",
        "pythia_metadata_json",
    ]

    questions: List[Dict[str, Any]] = []
    for r in rows:
        q = dict(zip(cols, r))
        try:
            q["scenario_ids"] = json.loads(q["scenario_ids_json"] or "[]")
        except Exception:
            q["scenario_ids"] = []
        try:
            q["metadata"] = json.loads(q["pythia_metadata_json"] or "{}")
        except Exception:
            q["metadata"] = {}
        questions.append(q)
    return questions


def resolve_run_id_for_hs(
    con: duckdb.DuckDBPyConnection,
    hs_run_id: str,
    explicit_run_id: Optional[str] = None,
) -> Optional[str]:
    """
    If explicit_run_id is given, return it; otherwise pick the latest run_id
    in forecasts_ensemble for any question belonging to hs_run_id.
    """

    if explicit_run_id:
        return explicit_run_id

    row = con.execute(
        """
        SELECT fe.run_id
        FROM forecasts_ensemble fe
        JOIN questions q
          ON fe.question_id = q.question_id
        WHERE q.hs_run_id = ?
        ORDER BY fe.created_at DESC
        LIMIT 1
        """,
        [hs_run_id],
    ).fetchone()
    return row[0] if row else None


def load_spd_ev_summary_for_questions(
    con: duckdb.DuckDBPyConnection,
    question_ids: List[str],
    run_id: str,
) -> Dict[str, Dict[int, float]]:
    """
    Return {question_id: {month_index: ev_value}} using forecasts_ensemble.

    We assume ev_value is stored per (run_id, question_id, month_index, bucket_index)
    and either repeated or only present for one bucket; we just pick the first non-null.
    """

    if not question_ids:
        return {}
    placeholders = ",".join(["?"] * len(question_ids))

    rows = con.execute(
        f"""
        SELECT question_id, month_index, ev_value
        FROM forecasts_ensemble
        WHERE run_id = ? AND question_id IN ({placeholders})
          AND ev_value IS NOT NULL
        ORDER BY question_id, month_index
        """,
        [run_id, *question_ids],
    ).fetchall()

    result: Dict[str, Dict[int, float]] = {}
    for question_id, month_idx, ev_val in rows:
        qid = str(question_id)
        month_idx_int = int(month_idx)
        if qid not in result:
            result[qid] = {}
        if month_idx_int not in result[qid]:
            result[qid][month_idx_int] = float(ev_val)
    return result


def load_calibration_summary_for_questions(
    con: duckdb.DuckDBPyConnection,
    questions: List[Dict[str, Any]],
) -> Dict[Tuple[str, str], Dict[str, Any]]:
    """
    Return mapping (hazard_code, metric) -> {weights: [...], advice_text: str}
    """

    keys = sorted({((q.get("hazard_code") or "").upper(), q.get("metric") or "") for q in questions})
    result: Dict[Tuple[str, str], Dict[str, Any]] = {}

    for hz, metric in keys:
        if not hz or not metric:
            continue
        w_rows = con.execute(
            """
            SELECT model_name, weight
            FROM calibration_weights
            WHERE hazard_code = ? AND metric = ?
            ORDER BY model_name
            """,
            [hz, metric],
        ).fetchall()
        weights = [{"model_name": r[0], "weight": float(r[1])} for r in w_rows]

        advice_row = con.execute(
            """
            SELECT advice_text
            FROM calibration_advice
            WHERE hazard_code = ? AND metric = ?
            ORDER BY updated_at DESC
            LIMIT 1
            """,
            [hz, metric],
        ).fetchone()
        advice_text = advice_row[0] if advice_row else ""

        result[(hz, metric)] = {"weights": weights, "advice_text": advice_text}
    return result


def build_country_index(
    hs_scenarios: List[Dict[str, Any]],
    questions: List[Dict[str, Any]],
    ev_summary: Dict[str, Dict[int, float]],
    calib_summary: Dict[Tuple[str, str], Dict[str, Any]],
) -> Dict[str, Dict[str, Any]]:
    """
    Build a structure:
      {iso3: {
          "scenarios": [...],
          "questions": [...]
      }}
    Where each question includes EVs and a one-line calibration note.
    """

    countries: Dict[str, Dict[str, Any]] = defaultdict(lambda: {"scenarios": [], "questions": []})

    for s in hs_scenarios:
        iso3 = (s.get("iso3") or "").upper()
        countries[iso3]["scenarios"].append(s)

    for q in questions:
        iso3 = (q.get("iso3") or "").upper()
        hazard = (q.get("hazard_code") or "").upper()
        metric = q.get("metric") or ""
        qid = q["question_id"]
        evs = ev_summary.get(qid, {})
        calib = calib_summary.get((hazard, metric), {"weights": [], "advice_text": ""})

        weights = calib.get("weights") or []
        note = ""
        if weights:
            sorted_w = sorted(weights, key=lambda w: w["weight"], reverse=True)
            top = sorted_w[0]
            note = f"Top weight: {top['model_name']} ({top['weight']:.2f})"
        elif calib.get("advice_text"):
            note = calib["advice_text"][:160]
        else:
            note = "No calibration weights recorded."

        q_compact = {
            "question_id": qid,
            "wording": q.get("wording") or "",
            "hazard_code": hazard,
            "metric": metric,
            "window_start_date": q.get("window_start_date") or "",
            "window_end_date": q.get("window_end_date") or "",
            "ev_by_month": evs,
            "calibration_note": note,
        }
        countries[iso3]["questions"].append(q_compact)

    return countries


def build_pythia_run_docx(
    out_path: Path,
    hs_run: Dict[str, Any],
    countries: Dict[str, Dict[str, Any]],
    run_id: Optional[str],
) -> None:
    doc = Document()

    hs_run_id = hs_run["hs_run_id"]
    generated_at = hs_run.get("generated_at")
    config_profile = hs_run.get("config_profile") or ""
    git_sha = hs_run.get("git_sha") or ""
    countries_json = hs_run.get("countries_json") or "[]"
    try:
        hs_countries = json.loads(countries_json)
    except Exception:
        hs_countries = []

    doc.add_heading("Pythia HS + Forecaster Run Report", level=0)
    doc.add_paragraph(f"HS Run ID: {hs_run_id}")
    if run_id:
        doc.add_paragraph(f"Forecaster Run ID: {run_id}")
    if generated_at:
        doc.add_paragraph(f"Generated at: {generated_at}")
    if config_profile:
        doc.add_paragraph(f"Forecaster config profile: {config_profile}")
    if git_sha:
        doc.add_paragraph(f"Git SHA: {git_sha}")

    n_questions = sum(len(v["questions"]) for v in countries.values())
    doc.add_paragraph(f"Countries in HS run: {', '.join(sorted(set(hs_countries)))}")
    doc.add_paragraph(f"Total Pythia questions for this HS run: {n_questions}")

    for iso3 in sorted(countries.keys()):
        country_block = countries[iso3]
        scenarios = country_block["scenarios"]
        qs = country_block["questions"]

        doc.add_heading(f"Country: {iso3}", level=1)

        doc.add_paragraph("Horizon Scanner scenarios:")

        if scenarios:
            table = doc.add_table(rows=1, cols=5)
            hdr = table.rows[0].cells
            hdr[0].text = "Scenario Title"
            hdr[1].text = "Hazard"
            hdr[2].text = "Likely Month"
            hdr[3].text = "Probability (%)"
            hdr[4].text = "PA Best Guess"

            for s in scenarios:
                row = table.add_row().cells
                row[0].text = str(s.get("scenario_title") or "")
                row[1].text = str(s.get("hazard_code") or "")
                row[2].text = str(s.get("likely_month") or "")
                row[3].text = str(s.get("probability_pct") or "")
                row[4].text = str(s.get("pa_best_guess") or "")
        else:
            doc.add_paragraph("No HS scenarios found for this country.")

        doc.add_paragraph()
        doc.add_paragraph("Pythia forecast questions:")

        if qs:
            for q in qs:
                doc.add_paragraph(f"• {q['wording']}", style="List Bullet")

                if q["window_start_date"] and q["window_end_date"]:
                    doc.add_paragraph(
                        f"  Window: {q['window_start_date']} → {q['window_end_date']}"
                    )

                if q["ev_by_month"]:
                    table_ev = doc.add_table(rows=1 + len(q["ev_by_month"]), cols=2)
                    hdr = table_ev.rows[0].cells
                    hdr[0].text = "Month"
                    hdr[1].text = "Expected PA"
                    for m in sorted(q["ev_by_month"].keys()):
                        row = table_ev.add_row().cells
                        row[0].text = str(m)
                        row[1].text = f"{float(q['ev_by_month'][m]):,.0f}"
                else:
                    doc.add_paragraph("  (No expected PA values recorded yet.)")

                calib_note = q.get("calibration_note") or ""
                if calib_note:
                    doc.add_paragraph(f"  Calibration: {calib_note}")
                doc.add_paragraph()
        else:
            doc.add_paragraph("No Pythia forecast questions for this HS run in this country.")

    doc.add_heading("Further Detail", level=1)
    doc.add_paragraph(
        "For detailed per-question HS context, research bundle, full SPD tables, "
        "and raw model explanations, see the individual question reports "
        "(e.g. files named 'Forecast_Q_{question_id}.docx')."
    )

    doc.save(str(out_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Export a Pythia HS+Forecaster run to Word.")
    parser.add_argument(
        "--hs-run-id",
        required=True,
        help="HS run id to export (e.g. hs_20251128T092046).",
    )
    parser.add_argument(
        "--run-id",
        required=False,
        help=(
            "Optional Forecaster run_id. If not provided, the latest run_id with "
            "forecasts_ensemble rows for this HS run will be used."
        ),
    )
    args = parser.parse_args()

    hs_run_id = args.hs_run_id
    explicit_run_id = args.run_id

    con = _connect_duckdb()
    try:
        hs_run, hs_scenarios = load_hs_run_summary(con, hs_run_id)
        questions = load_pythia_questions_for_hs(con, hs_run_id)

        if not questions:
            raise RuntimeError(f"No Pythia questions found for hs_run_id={hs_run_id}.")

        run_id = resolve_run_id_for_hs(con, hs_run_id, explicit_run_id)
        if not run_id:
            raise RuntimeError(
                f"No forecasts_ensemble rows found for hs_run_id={hs_run_id} "
                "and no --run-id was provided."
            )

        qids = [q["question_id"] for q in questions]
        ev_summary = load_spd_ev_summary_for_questions(con, qids, run_id)
        calib_summary = load_calibration_summary_for_questions(con, questions)
        countries = build_country_index(hs_scenarios, questions, ev_summary, calib_summary)
    finally:
        con.close()

    out_name = f"Pythia_Run_{hs_run_id}.docx"
    out_path = Path(out_name)
    build_pythia_run_docx(out_path, hs_run, countries, run_id)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
